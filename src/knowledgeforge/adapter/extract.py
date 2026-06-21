"""Text extractors — one per file type, all return plain text."""
from __future__ import annotations

import csv
import io
import json
from pathlib import Path
from typing import Any


def extract_text(path: Path) -> str:
    """Return plain text from any supported file. Empty string on failure."""
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _pdf(path)
        if suffix in {".docx"}:
            return _docx(path)
        if suffix in {".doc"}:
            return _doc_fallback(path)
        if suffix in {".html", ".htm"}:
            return _html(path)
        if suffix == ".csv":
            return _csv(path)
        if suffix == ".tsv":
            return _csv(path, delimiter="\t")
        if suffix in {".json", ".jsonl"}:
            return _json(path)
        if suffix in {".yaml", ".yml"}:
            return _yaml(path)
        # everything else: read as text (code, md, txt, rst, toml, xml, etc.)
        return path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return ""


# ── per-format extractors ──────────────────────────────────────────────────────

def _pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n\n".join(parts)


def _docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _doc_fallback(path: Path) -> str:
    # .doc (old binary Word) — try reading as text, mostly useful for extraction attempts
    return path.read_text(encoding="utf-8", errors="replace")


def _html(path: Path) -> str:
    from bs4 import BeautifulSoup, Tag
    raw = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "lxml")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    lines: list[str] = []
    for tag in soup.find_all(True):
        if not isinstance(tag, Tag):
            continue
        name = tag.name.lower()
        text = tag.get_text(strip=True)
        if not text:
            continue
        if name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = int(name[1])
            lines.append(f"{'#' * level} {text}")
        elif name == "a":
            href = tag.get("href", "")
            if href and text:
                lines.append(f"[{text}]({href})")
        elif name in {"p", "li", "td", "th", "blockquote", "pre", "code"}:
            lines.append(text)

    return "\n".join(lines) if lines else soup.get_text(separator="\n")


def _csv(path: Path, delimiter: str = ",") -> str:
    raw = path.read_text(encoding="utf-8", errors="replace")
    reader = csv.DictReader(io.StringIO(raw), delimiter=delimiter)
    lines: list[str] = []
    headers: list[str] = []
    for i, row in enumerate(reader):
        if i == 0:
            headers = list(row.keys())
            lines.append("Columns: " + ", ".join(headers))
        cells = [f"{k}: {v}" for k, v in row.items() if v and v.strip()]
        if cells:
            lines.append(" | ".join(cells))
        if i >= 500:  # cap rows to avoid giant context
            lines.append(f"... ({i+1}+ rows total)")
            break
    return "\n".join(lines)


def _json(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="replace")
    lines: list[str] = []
    if path.suffix == ".jsonl":
        for line in raw.splitlines()[:500]:
            line = line.strip()
            if line:
                try:
                    obj = json.loads(line)
                    lines.append(_flatten_json(obj))
                except json.JSONDecodeError:
                    lines.append(line)
    else:
        try:
            obj = json.loads(raw)
            lines.append(_flatten_json(obj))
        except json.JSONDecodeError:
            lines.append(raw[:10000])
    return "\n".join(lines)


def _flatten_json(obj: Any, prefix: str = "", depth: int = 0) -> str:
    if depth > 5:
        return ""
    parts: list[str] = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            key = f"{prefix}.{k}" if prefix else k
            if isinstance(v, (dict, list)):
                parts.append(_flatten_json(v, key, depth + 1))
            elif v is not None and str(v).strip():
                parts.append(f"{key}: {v}")
    elif isinstance(obj, list):
        for i, item in enumerate(obj[:50]):
            parts.append(_flatten_json(item, f"{prefix}[{i}]", depth + 1))
    else:
        if obj is not None:
            parts.append(f"{prefix}: {obj}")
    return "\n".join(p for p in parts if p)


def _yaml(path: Path) -> str:
    try:
        import yaml
        raw = path.read_text(encoding="utf-8", errors="replace")
        obj = yaml.safe_load(raw)
        return _flatten_json(obj) if obj else raw
    except Exception:
        return path.read_text(encoding="utf-8", errors="replace")
